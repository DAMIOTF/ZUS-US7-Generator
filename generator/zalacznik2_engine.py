# -*- coding: utf-8 -*-
"""
Silnik generowania Załącznika nr 2 (DOCX).

Publiczne API:
    fill_zalacznik2(person, date_str, options) -> bytes
"""
import io

from docx import Document

from .constants import resource_path

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

TEMPLATE_DOCX = resource_path("zalacznik2/szablon_pusty.docx")


# ─── Akceptuj ślubione zmiany redaktora ──────────────────────────────────────

def _accept_all_tracked_changes(doc: Document) -> None:
    """Akceptuje wszystkie śledzone zmiany w dokumencie."""
    body = doc.element.body

    for del_elem in body.findall(".//{%s}del" % W):
        parent = del_elem.getparent()
        if parent is not None:
            parent.remove(del_elem)

    for ins_elem in body.findall(".//{%s}ins" % W):
        parent = ins_elem.getparent()
        if parent is None:
            continue
        idx = list(parent).index(ins_elem)
        for child in list(ins_elem):
            parent.insert(idx, child)
            idx += 1
        parent.remove(ins_elem)


# ─── Adres ────────────────────────────────────────────────────────────────────

def _build_address(person: dict) -> str:
    """Buduje adres z danych osoby.

    Gdy poczta != miejscowosc, miejscowosc to wioska (trafia jako lokalizacja),
    a poczta jest właściwym miastem.
    """
    ulica       = person.get("ulica", "").strip()
    miejscowosc = person.get("miejscowosc", "").strip()
    poczta      = person.get("poczta", "").strip()
    nr_domu     = person.get("nr_domu", "").strip()
    nr_lokalu   = person.get("nr_lokalu", "").strip()
    kod         = person.get("kod_pocztowy", "").strip()

    # Rozpoznanie wsi
    if poczta and poczta != miejscowosc and not ulica:
        ulica = miejscowosc
    miasto = poczta or miejscowosc

    nr = nr_domu
    if nr_domu and nr_lokalu:
        nr = f"{nr_domu}/{nr_lokalu}"

    parts = []
    if ulica:
        parts.append(ulica)
    if nr:
        parts.append(nr)
    if kod:
        parts.append(kod)
    if miasto:
        parts.append(miasto)
    return ", ".join(parts)


# ─── Generator DOCX ──────────────────────────────────────────────────────────

def fill_zalacznik2(person: dict, date_str: str, options: dict) -> bytes:
    """Wypełnia szablon Załącznika nr 2 i zwraca bajty pliku DOCX."""
    imie = person.get("imie", "")
    nazwisko = person.get("nazwisko", "")
    pesel = person.get("pesel", "")
    telefon = person.get("telefon", "")
    adres = _build_address(person)

    if pesel.isdigit() and len(pesel) < 11:
        pesel = pesel.zfill(11)

    doc = Document(TEMPLATE_DOCX)
    _accept_all_tracked_changes(doc)

    # Para 4: "Ja niżej podpisany/a <IMIĘ NAZWISKO>"
    para4 = doc.paragraphs[4]
    run = para4.add_run(f"{imie} {nazwisko}")
    run.bold = False
    run.underline = False

    # Para 5: "zamieszkały/a <ADRES>"
    para5 = doc.paragraphs[5]
    run = para5.add_run(adres)
    run.bold = False
    run.underline = False

    # Para 6: "numer PESEL <PESEL>"
    para6 = doc.paragraphs[6]
    run = para6.add_run(pesel)
    run.bold = False
    run.underline = False

    # Para 7, run 2: numer telefonu
    para7 = doc.paragraphs[7]
    if len(para7.runs) > 2:
        para7.runs[2].text = f"{telefon} "

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
