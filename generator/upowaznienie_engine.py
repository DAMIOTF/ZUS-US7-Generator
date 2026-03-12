# -*- coding: utf-8 -*-
"""
Silnik generowania Upoważnienia ZUS (DOCX).

Publiczne API:
    fill_upowaznienie(person, date_str, options) -> bytes
"""
import io

from docx import Document
from docx.shared import Pt, RGBColor

from .constants import resource_path

TEMPLATE_DOCX = resource_path("UPOWAŻNIENIE ZUS.docx")


def _build_address(person: dict) -> str:
    """Buduje pełny adres z danych osoby.

    Gdy poczta != miejscowosc, miejscowosc to wioska (trafia jako lokalizacja),
    a poczta jest właściwym miastem.
    """
    ulica       = person.get("ulica", "").strip()
    miejscowosc = person.get("miejscowosc", "").strip()
    poczta      = person.get("poczta", "").strip()
    nr_domu     = person.get("nr_domu", "").strip()
    nr_lokalu   = person.get("nr_lokalu", "").strip()
    kod         = person.get("kod_pocztowy", "").strip()

    # Rozpoznanie wsi
    if poczta and poczta != miejscowosc and not ulica:
        ulica = miejscowosc
    miasto = poczta or miejscowosc

    nr = nr_domu
    if nr_domu and nr_lokalu:
        nr = f"{nr_domu}/{nr_lokalu}"

    parts = []
    if ulica:
        parts.append(ulica)
    if nr:
        parts.append(nr)
    if kod:
        parts.append(kod)
    if miasto:
        parts.append(miasto)
    return ", ".join(parts)


def fill_upowaznienie(person: dict, date_str: str, options: dict) -> bytes:
    """Wypełnia szablon Upoważnienia ZUS i zwraca bajty pliku DOCX."""
    imie = person.get("imie", "")
    nazwisko = person.get("nazwisko", "")
    pesel = person.get("pesel", "")
    adres = _build_address(person)

    if pesel.isdigit() and len(pesel) < 11:
        pesel = pesel.zfill(11)

    doc = Document(TEMPLATE_DOCX)

    BLACK = RGBColor(0, 0, 0)

    def _set_run(run, text):
        """Ustawia tekst runa i wymusza czarny kolor bez pogrubienia."""
        run.text = text
        run.bold = False
        run.font.color.rgb = BLACK

    # P[0] R[1]: DATA → data wniosku
    para0 = doc.paragraphs[0]
    if len(para0.runs) > 1:
        _set_run(para0.runs[1], date_str)

    # P[3] R[1]: IMIĘ NAZWISKO → imię i nazwisko
    para3 = doc.paragraphs[3]
    if len(para3.runs) > 1:
        _set_run(para3.runs[1], f"{imie} {nazwisko}")

    # P[4] R[1]: PESEL → numer PESEL
    para4 = doc.paragraphs[4]
    if len(para4.runs) > 1:
        _set_run(para4.runs[1], pesel)

    # P[5] R[5]: adres zamieszkania
    para5 = doc.paragraphs[5]
    if len(para5.runs) > 5:
        _set_run(para5.runs[5], adres)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
