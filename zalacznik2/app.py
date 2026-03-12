import openpyxl
import shutil
import os
from docx import Document
from lxml import etree

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# ─────────────────────────── Akceptuj zmiany redaktora ─────────

def accept_all_tracked_changes(doc):
    """
    Akceptuje wszystkie śledzone zmiany w dokumencie:
    - usuwa <w:del> (odrzucone treści – zaakceptowanie = usunięcie)
    - odwija <w:ins> (wstawione treści – zaakceptowanie = zachowanie dzieci)
    """
    body = doc.element.body

    # Usuń w:del (włącznie z w:del wewnątrz w:rPr – marker akapitu)
    for del_elem in body.findall('.//{%s}del' % W):
        parent = del_elem.getparent()
        if parent is not None:
            parent.remove(del_elem)

    # Odwiń w:ins – zachowaj dzieci, usuń opakowanie
    for ins_elem in body.findall('.//{%s}ins' % W):
        parent = ins_elem.getparent()
        if parent is None:
            continue
        idx = list(parent).index(ins_elem)
        for child in list(ins_elem):
            parent.insert(idx, child)
            idx += 1
        parent.remove(ins_elem)


# ─────────────────────────── Ścieżki ───────────────────────────
BASE_DIR      = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(BASE_DIR, 'szablon pusty.docx')
EXCEL_PATH    = os.path.join(BASE_DIR, 'baza.xlsx')
OUTPUT_DIR    = os.path.join(BASE_DIR, 'wyniki')

os.makedirs(OUTPUT_DIR, exist_ok=True)


# ─────────────────────────── Pomocnicze ────────────────────────

def clean(val):
    """Zwróć czysty string z dowolnej wartości komórki."""
    if val is None:
        return ''
    if isinstance(val, float):
        return str(int(val)) if val.is_integer() else str(val)
    return str(val).strip()


def build_address(row):
    """
    Buduje adres w formacie:
      {kod} {poczta} [{miejscowość jeśli inna niż poczta}] [ul. {ulica}] {nr_domu}[/{nr_lokalu}]
    Przykład z szablonu: 68-100 Żagań Pożarów 30A
    """
    kod         = clean(row.get('Kod pocztowy', ''))
    poczta      = clean(row.get('Poczta', ''))
    miejscowosc = clean(row.get('Miejscowość ', ''))   # nagłówek ma spację na końcu
    ulica       = clean(row.get('Ulica', ''))
    nr_domu     = clean(row.get('Numer domu', ''))
    nr_lokalu   = clean(row.get('Numer Lokalu', ''))

    parts = []
    if kod:
        parts.append(kod)
    if poczta:
        parts.append(poczta)
    # jeśli miejscowość różni się od poczty, dodaj ją (unikaj powtórzenia)
    if miejscowosc and miejscowosc != poczta:
        parts.append(miejscowosc)
    if ulica:
        parts.append(f'ul. {ulica}')
    if nr_domu:
        addr_num = nr_domu
        if nr_lokalu:
            addr_num += f'/{nr_lokalu}'
        parts.append(addr_num)

    return ' '.join(parts)


def fill_document(row):
    """
    Kopiuje szablon i uzupełnia go danymi jednej osoby.
    Zwraca nazwę pliku wynikowego.
    """
    imie     = clean(row.get('Imię (Imiona)', ''))
    nazwisko = clean(row.get('Nazwisko ', ''))   # nagłówek ma spację na końcu
    pesel    = clean(row.get('PESEL', ''))
    telefon  = clean(row.get('Telefon', ''))
    adres    = build_address(row)

    # Zerowanie do 11 znaków, gdyby PESEL trafił jako int (bez wiodącego zera)
    if pesel.isdigit() and len(pesel) < 11:
        pesel = pesel.zfill(11)

    # Nazwa pliku: Imię_Imię2_Nazwisko.docx
    safe_name   = f"{imie.replace(' ', '_')}_{nazwisko}.docx"
    output_path = os.path.join(OUTPUT_DIR, safe_name)

    # Kopiuj szablon
    shutil.copy(TEMPLATE_PATH, output_path)

    # Otwórz kopię, przyjmij wszystkie śledzone zmiany, potem uzupełnij
    doc = Document(output_path)
    accept_all_tracked_changes(doc)

    # --- Para 4: "Ja niżej podpisany/a <IMIĘ NAZWISKO>" ---
    para4 = doc.paragraphs[4]
    run = para4.add_run(f'{imie} {nazwisko}')
    run.bold      = False
    run.underline = False

    # --- Para 5: "zamieszkały/a  <ADRES>" ---
    para5 = doc.paragraphs[5]
    run = para5.add_run(adres)
    run.bold      = False
    run.underline = False

    # --- Para 6: "numer PESEL <PESEL>" ---
    para6 = doc.paragraphs[6]
    run = para6.add_run(pesel)
    run.bold      = False
    run.underline = False

    # --- Para 7, run 2: zastąp numer telefonu ---
    # Szablon ma tu zakodowany '663158966 ' – zastępujemy właściwym numerem
    para7 = doc.paragraphs[7]
    if len(para7.runs) > 2:
        para7.runs[2].text = f'{telefon} '

    doc.save(output_path)
    return safe_name


# ─────────────────────────── Główna logika ─────────────────────

def main():
    print(f'Odczyt bazy: {EXCEL_PATH}')
    wb = openpyxl.load_workbook(EXCEL_PATH)
    ws = wb.active

    # Nagłówki (wiersz 1)
    headers = [cell.value for cell in next(ws.iter_rows(min_row=1, max_row=1))]
    print(f'Nagłówki: {headers}')
    print(f'Wierszy z danymi: {ws.max_row - 1}\n')

    print('Generowanie dokumentów...')
    count  = 0
    errors = []

    for row_vals in ws.iter_rows(min_row=2, max_row=ws.max_row, values_only=True):
        if row_vals[0] is None:
            continue
        row = dict(zip(headers, row_vals))
        try:
            filename = fill_document(row)
            print(f'  ✓  {filename}')
            count += 1
        except Exception as e:
            imie     = clean(row.get('Imię (Imiona)', '?'))
            nazwisko = clean(row.get('Nazwisko ', '?'))
            msg = f'BŁĄD dla {imie} {nazwisko}: {e}'
            print(f'  ✗  {msg}')
            errors.append(msg)

    print(f'\nGotowe! Utworzono {count} plików w: {OUTPUT_DIR}')
    if errors:
        print(f'\nBłędy ({len(errors)}):')
        for e in errors:
            print(f'  - {e}')


if __name__ == '__main__':
    main()